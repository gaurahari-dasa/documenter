from docx import Document
import re
import json


def replace_text_in_docx(doc: Document, placeholders: dict):
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if not value:
                # Construct regex pattern dynamically based on the tag
                pattern = re.compile(rf"\{{{re.escape(key)}: .*? :{re.escape(key)}\}}")
                para.text = re.sub(pattern, "", para.text)
            else:
                pattern = re.compile(
                    rf"\{{{re.escape(key)}: (.*?) :{re.escape(key)}\}}"
                )
                para.text = re.sub(pattern, r"\1", para.text)


def fill_placeholders(template_path, output_path):
    def sub(m: re.Match):
        try:
            return placeholders[m.group(1)]
        except Exception as ex:
            print("No match:", ex)
            return m.group(0)

    # Load the Word template
    doc = Document(template_path)

    # Replace placeholders with actual values
    for paragraph in doc.paragraphs:
        paragraph.text = re.sub(r"\{\{([a-zA-Z0-9 _\-]+)\}\}", sub, paragraph.text)
        # for placeholder, value in placeholders.items():
        #     if placeholder in paragraph.text:
        #         paragraph.text = paragraph.text.replace(placeholder, value)

    replace_text_in_docx(doc, placeholders)

    # Save the filled document
    doc.save(output_path)


# Load the placeholders and their values from a JSON file
with open("placeholders.json", "r") as f:
    placeholders = json.load(f)


# Paths to the template and output files
# template_path = input("Path to the template: ").strip("\"'")
template_path = "documentation_template.docx"
output_path = "filled_documentation.docx"

# Fill the placeholders in the template and save the result
fill_placeholders(template_path, output_path)

print(f"The document has been filled and saved as {output_path}.")
